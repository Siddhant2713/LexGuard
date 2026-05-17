import fitz  # PyMuPDF
import io
from typing import Tuple


def _try_ocr(page: fitz.Page) -> str:
    """OCR fallback using pytesseract for image-heavy pages."""
    try:
        import pytesseract
        from PIL import Image

        pix = page.get_pixmap(dpi=200)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        return pytesseract.image_to_string(img)
    except ImportError:
        return ""
    except Exception:
        return ""


def extract_pdf(pdf_bytes: bytes) -> Tuple[str, dict]:
    """
    Extract text from a PDF file.

    Returns:
        raw_text: Full document text, pages joined with double newline.
        page_map: {page_num (1-indexed): page_text}

    Raises:
        ValueError: If no extractable text found even after OCR fallback.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page_map: dict = {}
    pages: list[str] = []

    for i, page in enumerate(doc):
        text = page.get_text("text").strip()

        # OCR fallback: if PyMuPDF finds fewer than 50 chars on this page,
        # it's likely an image-based scan — try tesseract
        if len(text) < 50:
            ocr_text = _try_ocr(page)
            if len(ocr_text.strip()) > len(text):
                text = ocr_text.strip()

        page_map[i + 1] = text
        pages.append(text)

    raw_text = "\n\n".join(pages)

    if len(raw_text.strip()) < 100:
        raise ValueError(
            "Could not extract readable text from this PDF. "
            "The file may be a scanned image without OCR data."
        )

    return raw_text, page_map


def extract_docx(docx_bytes: bytes) -> Tuple[str, dict]:
    """
    Extract text from a DOCX file.

    Returns:
        raw_text: Full document text.
        page_map: {1: full_text} — DOCX has no page concept, returns single entry.
    """
    try:
        from docx import Document
        import io as _io

        doc = Document(_io.BytesIO(docx_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        raw_text = "\n\n".join(paragraphs)

        if len(raw_text.strip()) < 100:
            raise ValueError("DOCX file appears to be empty or contains no readable text.")

        return raw_text, {1: raw_text}

    except ImportError:
        raise ValueError("python-docx is not installed. Cannot parse DOCX files.")


def extract_document(file_bytes: bytes, filename: str) -> Tuple[str, dict]:
    """
    Universal document extractor. Dispatches to PDF or DOCX parser based on filename.
    """
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_pdf(file_bytes)
    elif lower.endswith(".docx") or lower.endswith(".doc"):
        return extract_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Only PDF and DOCX are supported.")

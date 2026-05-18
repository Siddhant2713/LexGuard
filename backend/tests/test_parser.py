import pytest
from io import BytesIO
from parser import extract_document, extract_docx


def test_unsupported_extension():
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_document(b"dummy content", "contract.xlsx")


def test_docx_empty_raises():
    """Empty DOCX (valid structure, no content) should raise ValueError."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    buf = BytesIO()
    doc.save(buf)
    with pytest.raises(ValueError, match="empty"):
        extract_docx(buf.getvalue())


def test_docx_with_content():
    """DOCX with paragraphs should extract text correctly."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_paragraph("This is an employment agreement between Acme Corp and John Doe.")
    doc.add_paragraph("The non-compete clause restricts work for 24 months after termination.")
    buf = BytesIO()
    doc.save(buf)

    raw_text, page_map = extract_docx(buf.getvalue())
    assert "employment agreement" in raw_text.lower()
    assert "non-compete" in raw_text.lower()
    assert len(page_map) == 1
